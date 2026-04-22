# -*- coding: utf-8 -*-
"""
双周报 Word 生成：从 Excel 需求台账按 Title_Map 映射，追加写入双周报 Word 模板表格。

参考模式：merge_requirement_to_template-V10.py（配置、日志、路径预检、utf-8-sig、PyInstaller 基目录）。
20260325：对需求的表头做处理，可以增加【新增】，并在后面加上需求复杂度。
"""
from __future__ import annotations

import configparser
import difflib
import json
import sys
import uuid
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- 修订说明（编号在代码中对应引用）----------
# [WR-001] 使用与程序同目录的 workreport.ini，可配置 TEMPLATE_PATH、INPUT_PATHS、Title_Map。
# [WR-002] INPUT_PATHS 支持多文件（换行/逗号/分号分隔），数据合并后再写入 Word。
# [WR-003] 新数据追加在模板表格已有数据行之后；按「受理时间」升序（YYYY-MM-DD 解析与排序）。
# [WR-004] 「问题」列新行统一填「暂无」；列映射由 Title_Map（JSON）定义。
# [WR-005] 日志文件名为「脚本名_日期时间.log」，记录新增条数等，并同步输出到标准输出。
# [WR-006] 路径存在性预检查与友好报错建议（模板 .docx、输入 .xlsx）。
# [WR-007] 首条新增行插入书签 NewImportedData，便于 Word 中「转到」书签定位（标准 docx 无法在打开时自动将光标跳到该书签）。
# [WR-008] 配置文件使用 utf-8-sig 读取，兼容 Windows 记事本「UTF-8」含 BOM。
# [WR-009] PyInstaller 打包后从可执行文件所在目录读取 workreport.ini。
# [WR-010] INPUT_ROW_LIMIT：0 或留空表示不限制行数；正整数表示仅处理合并后的前 N 行（便于试跑）。
# [WR-011] 多源 Excel 列映射到同一 Word 列时，仅非空值覆盖，避免后序空值清空先序已填（如「创建时间」与「回复时间/ITM回函时间」同映「受理时间」）。
# [WR-012] 新增行通过复制模板中「样式来源行」的 tr 结构追加，使新行与模板表格格式一致。
# [WR-013] 写入单元格时尽量保留克隆行原有字体等格式，不强制改为宋体。
# [WR-014] Title_Map 支持 JSON 值为字符串或字符串数组（一对多映射）。
# [WR-015] 可选 WORD_TABLE_INDEX 指定使用文档中第几张表（从 0 开始），默认 0。
# [WR-016] 可选 Compose_Map（JSON）：按 Word 表头列名拼接多列与固定字符串，表达式形如
#          "固定"+Excel列名+"\n"+另一列；双引号内支持 \\n \\t \\r \\\\ \\" 等转义；在 Title_Map 之后生效并覆盖同名列。
# [WR-017]  修改日期格式为  YYYY/MM/DD，修改“需求复杂度”的格式为“（实施复杂度：标准）”，选用的列也采用“实施复杂度”
# [WR-018]  此处修改，解决了由于 word 表头有“受理时间”改为“需求审批\n时间”造成的日期格式显示没有按照yyyy/mm/dd 显示的问题，
#       不修改将按照excel表格原始格式显示

CONFIG_FILENAME = "workreport-JK.ini"
CONFIG_SECTION = "PATHS"
DEFAULT_ISSUE_TEXT = "暂无"
BOOKMARK_NEW_DATA = "NewImportedData"


def normalize(s: object) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    if isinstance(s, str) and s.strip() == "":
        return ""
    try:
        if pd.isna(s):
            return ""
    except (TypeError, ValueError):
        pass
    return str(s).strip()


def get_base_dir() -> Path:
    # [WR-009] 打包后使用 exe 所在目录，开发时使用脚本目录
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).resolve().parent


def build_output_path(template_path: Path) -> Path:
    # 模板名 + 生成日期；若同日重复运行则自动追加序号避免覆盖
    date_part = datetime.now().strftime("%Y%m%d")
    base = template_path.with_name(f"{template_path.stem}_{date_part}{template_path.suffix}")
    if not base.exists():
        return base
    n = 2
    while True:
        candidate = template_path.with_name(f"{template_path.stem}_{date_part}_{n}{template_path.suffix}")
        if not candidate.exists():
            return candidate
        n += 1


def build_log_path(script_path: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return script_path.with_name(f"{script_path.stem}_{ts}.log")


def parse_input_paths(raw_value: str) -> list[Path]:
    # [WR-002] 多路径解析（与 V10 一致）
    normalized = raw_value.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("，", ",").replace("；", ";")
    parts: list[str] = []
    for line in normalized.split("\n"):
        for block in line.split(";"):
            for piece in block.split(","):
                candidate = piece.strip().strip('"').strip("'")
                if candidate:
                    parts.append(candidate)
    return [Path(p) for p in parts]


def normalize_filename_key(name: str) -> str:
    return (
        name.strip()
        .replace("（", "(")
        .replace("）", ")")
        .replace("【", "[")
        .replace("】", "]")
        .replace("，", ",")
        .replace("。", ".")
        .replace("：", ":")
        .replace(" ", "")
        .lower()
    )


def resolve_similar_path(path: Path, suffixes: tuple[str, ...]) -> Path:
    if path.exists():
        return path
    parent = path.parent
    if not parent.exists():
        return path
    target_key = normalize_filename_key(path.name)
    candidates: list[Path] = []
    for suf in suffixes:
        candidates.extend(parent.glob(f"*{suf}"))
    for candidate in candidates:
        if normalize_filename_key(candidate.name) == target_key:
            return candidate
    return path


def build_path_hint(path: Path, kind: str) -> str:
    # [WR-006] 路径不存在时的友好提示
    parent = path.parent
    if not parent.exists():
        return (
            f"路径不存在: {path}\n"
            f"建议: 确认目录是否存在 -> {parent}\n"
            "建议: 检查盘符、中文括号与全半角是否写错。"
        )
    if kind == "excel":
        names = [p.name for p in parent.iterdir() if p.is_file() and p.suffix.lower() in {".xlsx", ".xlsm", ".xls"}]
    else:
        names = [p.name for p in parent.iterdir() if p.is_file() and p.suffix.lower() == ".docx"]
    if not names:
        return (
            f"路径不存在: {path}\n"
            f"建议: 目录下未发现符合条件的文件 -> {parent}\n"
        )
    matches = difflib.get_close_matches(path.name, names, n=3, cutoff=0.45)
    if matches:
        cand = "\n".join(f"  - {m}" for m in matches)
        return f"路径不存在: {path}\n建议: 同目录下相似文件名：\n{cand}"
    sample = "\n".join(f"  - {n}" for n in names[:8])
    return f"路径不存在: {path}\n建议: 同目录下部分文件：\n{sample}"


def parse_title_map(raw_value: str) -> dict[str, list[str]]:
    # [WR-001][WR-014] Excel 列名 -> 一个或多个 Word 表头列名
    if not raw_value.strip():
        raise ValueError("配置项 Title_Map 不能为空")
    try:
        parsed = json.loads(raw_value)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Title_Map 不是合法 JSON: {exc}") from exc
    if not isinstance(parsed, dict):
        raise ValueError("Title_Map 必须是 JSON 对象")
    out: dict[str, list[str]] = {}
    for src, tgt in parsed.items():
        sk = normalize(src)
        if not sk:
            continue
        if isinstance(tgt, str):
            targets = [normalize(tgt)]
        elif isinstance(tgt, list):
            targets = [normalize(x) for x in tgt if isinstance(x, str)]
        else:
            raise ValueError(f"Title_Map 中 {src!r} 的值须为字符串或字符串数组")
        targets = [t for t in targets if t]
        if targets:
            out[sk] = targets
    if not out:
        raise ValueError("Title_Map 解析后为空")
    return out


def parse_compose_map(raw_value: str) -> dict[str, str]:
    # [WR-016] Word 表头列名 -> 拼接表达式（可选）
    print(f"compose_map: {raw_value}")
    if not raw_value.strip():
        return {}
    try:
        parsed = json.loads(raw_value)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Compose_Map 不是合法 JSON: {exc}") from exc
    if not isinstance(parsed, dict):
        raise ValueError("Compose_Map 必须是 JSON 对象")
    out: dict[str, str] = {}
    for word_col, expr in parsed.items():
        wk = normalize(word_col)
        if not wk:
            continue
        if not isinstance(expr, str):
            raise ValueError(f"Compose_Map 中 {word_col!r} 的值须为字符串（拼接表达式）")
        ex = expr.strip()
        if ex:
            out[wk] = ex
    return out


def unescape_compose_literal(s: str) -> str:
    # 双引号内字面量：\\n \\t \\r \\\\ \\" 等
    out: list[str] = []
    i = 0
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s):
            c = s[i + 1]
            if c == "n":
                out.append("\n")
            elif c == "r":
                out.append("\r")
            elif c == "t":
                out.append("\t")
            elif c == "\\":
                out.append("\\")
            elif c == '"':
                out.append('"')
            else:
                out.append(s[i : i + 2])
            i += 2
        else:
            out.append(s[i])
            i += 1
    return "".join(out)


def tokenize_compose_expression(expr: str) -> list[tuple[str, str]]:
    """将表达式拆为 ('lit', 文本) 或 ('col', Excel 列名)。"""
    s = expr.strip()
    i = 0
    parts: list[tuple[str, str]] = []

    def skip_ws() -> None:
        nonlocal i
        while i < len(s) and s[i].isspace():
            i += 1

    while True:
        skip_ws()
        if i >= len(s):
            break
        if s[i] == '"':
            i += 1
            buf: list[str] = []
            while i < len(s):
                if s[i] == "\\" and i + 1 < len(s):
                    buf.append(s[i : i + 2])
                    i += 2
                elif s[i] == '"':
                    i += 1
                    break
                else:
                    buf.append(s[i])
                    i += 1
            else:
                raise ValueError(f"拼接表达式中未闭合的双引号: {expr!r}")
            parts.append(("lit", unescape_compose_literal("".join(buf))))
        else:
            start = i
            while i < len(s) and s[i] != "+":
                i += 1
            col = s[start:i].strip()
            if col:
                parts.append(("col", col))
        skip_ws()
        if i >= len(s):
            break
        if s[i] != "+":
            raise ValueError(f"拼接表达式在位置 {i} 处应为 '+' 连接: {expr!r}")
        i += 1

    if not parts:
        raise ValueError(f"拼接表达式为空: {expr!r}")
    return parts


def compose_row_value(
    expr: str,
    row: object,
    *,
    title_map: dict[str, list[str]],
    accept_word_col: str | None,
    excel_columns: list[str],
) -> str:
    """按表达式从当前行取值并拼接。"""
    tokens = tokenize_compose_expression(expr)
    pieces: list[str] = []
    for kind, val in tokens:
        if kind == "lit":
            pieces.append(val)
            continue
        if val not in excel_columns:
            pieces.append("")
            continue
        raw = row[val]
        if accept_word_col and accept_word_col in title_map.get(val, []):
            pieces.append(format_accept_date(raw))
        else:
            pieces.append(normalize(raw))
    return "".join(pieces)


def read_input_excel(path: Path) -> pd.DataFrame:
    xls = pd.ExcelFile(path)
    preferred = "Sheet0"
    sheet_name = preferred if preferred in xls.sheet_names else xls.sheet_names[0]
    return pd.read_excel(path, sheet_name=sheet_name)


def parse_sort_date(value: object) -> datetime | None:
    # [WR-003] 按「受理时间」排序用
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    ts = pd.to_datetime(value, errors="coerce")
    if pd.isna(ts):
        return None
    return ts.to_pydatetime().replace(hour=0, minute=0, second=0, microsecond=0)


def format_accept_date(value: object) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    ts = pd.to_datetime(value, errors="coerce")
    if pd.isna(ts):
        return normalize(value)
    # [WR-017] 修改了日期格式 
    return ts.strftime("%Y/%m/%d")


def emit_log(lines: list[str], log_path: Path) -> None:
    # [WR-005]
    text = "\n".join(lines) + "\n"
    print(text, end="")
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with log_path.open("w", encoding="utf-8") as f:
        f.write(text)


def _cell_first_paragraph(cell):
    if not cell.paragraphs:
        return cell.add_paragraph()
    return cell.paragraphs[0]


def set_cell_text_preserve_style(cell, text: str) -> None:
    # [WR-013] 尽量保留克隆行原有 rPr，仅改首 run 文本并清空多余 run
    p = _cell_first_paragraph(cell)
    runs = p.runs
    while len(runs) > 1:
        p._p.remove(runs[-1]._r)
        runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text


def set_cell_text_preserve_style_with_bookmark(
    cell,
    text: str,
    *,
    bookmark_name: str,
    bookmark_id: str,
) -> None:
    # [WR-007][WR-013] 在首 run 两侧插入书签，保留该 run 的样式
    p = _cell_first_paragraph(cell)
    runs = p.runs
    while len(runs) > 1:
        p._p.remove(runs[-1]._r)
        runs = p.runs
    if not runs:
        p.add_run(text)
        runs = p.runs
    runs[0].text = text
    r0 = runs[0]._r
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    r0.addprevious(start)
    r0.addnext(end)


def append_cloned_row(table, template_tr) -> object:
    # [WR-012] 复制模板行 XML 追加到表格末尾，返回新行对象
    new_tr = deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def load_config(
    config_path: Path,
) -> tuple[Path, list[Path], list[Path], dict[str, list[str]], dict[str, str], int | None, int]:
    # [WR-001][WR-008][WR-009][WR-010][WR-015][WR-016]
    if not config_path.exists():
        raise FileNotFoundError(
            f"未找到配置文件: {config_path}\n"
            f"请在程序同目录创建 {CONFIG_FILENAME}，并配置 [{CONFIG_SECTION}] 节。"
        )
    parser = configparser.ConfigParser()
    parser.read(config_path, encoding="utf-8-sig")

    if CONFIG_SECTION not in parser:
        raise ValueError(f"配置文件缺少节: [{CONFIG_SECTION}]")

    section = parser[CONFIG_SECTION]
    template_raw = section.get("TEMPLATE_PATH", "").strip()
    inputs_raw = section.get("INPUT_PATHS", "").strip()
    title_map_raw = section.get("Title_Map", "").strip()
    compose_map_raw = section.get("Compose_Map", "").strip()
    # [WR-010] 默认不限制；正整数表示仅处理合并后的前 N 行
    limit_raw = section.get("INPUT_ROW_LIMIT", "0").strip()
    table_index_raw = section.get("WORD_TABLE_INDEX", "0").strip()

    if not template_raw:
        raise ValueError("TEMPLATE_PATH 不能为空")
    if not inputs_raw:
        raise ValueError("INPUT_PATHS 不能为空")

    template_path = resolve_similar_path(Path(template_raw), (".docx",))
    input_paths = [resolve_similar_path(p, (".xlsx", ".xlsm")) for p in parse_input_paths(inputs_raw)]

    if not template_path.exists():
        raise FileNotFoundError(
            "TEMPLATE_PATH 指向的 Word 模板不存在。\n" + build_path_hint(template_path, "docx")
        )

    missing = [p for p in input_paths if not p.exists()]
    existing = [p for p in input_paths if p.exists()]
    if not existing:
        hints = "\n\n".join(build_path_hint(p, "excel") for p in input_paths)
        raise FileNotFoundError("INPUT_PATHS 中没有可用 Excel 文件。\n" + hints)

    title_map = parse_title_map(title_map_raw)
    compose_map = parse_compose_map(compose_map_raw)

    row_limit: int | None = None
    try:
        n = int(limit_raw) if limit_raw else 0
    except ValueError:
        raise ValueError("INPUT_ROW_LIMIT 必须是整数（0 表示不限制）") from None
    if n > 0:
        row_limit = n

    try:
        table_index = int(table_index_raw) if table_index_raw else 0
    except ValueError:
        raise ValueError("WORD_TABLE_INDEX 必须是整数（从 0 开始）") from None

    return template_path, existing, missing, title_map, compose_map, row_limit, table_index


def main() -> None:
    base = get_base_dir()
    script_path = Path(sys.executable) if getattr(sys, "frozen", False) else Path(__file__).resolve()
    config_path = base / CONFIG_FILENAME

    template_path, input_paths, missing_inputs, title_map, compose_map, row_limit, table_index = load_config(
        config_path
    )
    if missing_inputs:
        print("警告: 以下输入文件不存在，已跳过：")
        for p in missing_inputs:
            print(f" - {p}")

    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("Word 模板中未找到表格，请检查模板。")
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f"WORD_TABLE_INDEX={table_index} 超出范围，当前文档共 {len(doc.tables)} 个表格。")

    table = doc.tables[table_index]
    num_cols = len(table.rows[0].cells)
    header_cells = [normalize(table.cell(0, c).text) for c in range(num_cols)]
    word_col_index = {name: i for i, name in enumerate(header_cells) if name}

    issue_col = "问题"
    if issue_col not in word_col_index:
        raise ValueError(f"Word 表格表头中未找到「{issue_col}」列，当前表头: {header_cells}")

    # [WR-012] 样式来源行：有数据时取最后一行，否则取表头行
    style_src_idx = len(table.rows) - 1 if len(table.rows) > 1 else 0
    template_tr = table.rows[style_src_idx]._tr

    original_data_rows = max(0, len(table.rows) - 1)

    parts: list[pd.DataFrame] = []
    for p in input_paths:
        df = read_input_excel(p)
        df.columns = [normalize(c) for c in df.columns]
        parts.append(df)

    all_in = pd.concat(parts, ignore_index=True)
    if row_limit is not None:
        # [WR-010]
        all_in = all_in.head(row_limit).reset_index(drop=True)
    # [WR-018]  此处修改，解决了由于 word 表头有“受理时间”改为“需求审批\n时间”造成的日期格式显示没有按照yyyy/mm/dd 显示的问题，
    # 不修改将按照excel表格原始格式显示
    accept_word_col = "需求审批\n时间" if "需求审批\n时间" in word_col_index else None
    if accept_word_col is None:
        for _src, tgts in title_map.items():
            for t in tgts:
                if "受理" in t and "时间" in t:
                    accept_word_col = t
                    break
            if accept_word_col:
                break

    new_rows: list[dict[str, str]] = []
    missing_targets: set[str] = set()
    excel_columns = list(all_in.columns)
    missing_compose_targets: set[str] = set()
    for word_col in compose_map:
        if word_col not in word_col_index:
            missing_compose_targets.add(word_col)

    for _, row in all_in.iterrows():
        rec: dict[str, str] = {h: "" for h in header_cells if h}
        for src_col, word_cols in title_map.items():
            for word_col in word_cols:
                if word_col not in word_col_index:
                    missing_targets.add(word_col)
                    continue
                if src_col not in all_in.columns:
                    continue
                raw = row[src_col]
                if accept_word_col and word_col == accept_word_col:
                    new_val = format_accept_date(raw)
                else:
                    new_val = normalize(raw)
                # [WR-011] 仅非空写入，多列映同一 Word 列时保留先出现的有效值
                if new_val:
                    rec[word_col] = new_val

        rec[issue_col] = DEFAULT_ISSUE_TEXT  # [WR-004]
        
        # 处理需求事项字段：添加【新增】前缀和需求复杂点
        requirement_col = "需求事项"
        if requirement_col in rec and rec[requirement_col]:
            original_text = rec[requirement_col]
            # 在前面添加【新增】
            modified_text = f"【新增】{original_text}"
            # 添加需求复杂点
            complexity_col = "实施复杂度"
            if complexity_col in all_in.columns:
                complexity_val = normalize(row[complexity_col])
                if complexity_val:
                    # [WR-017] 修改了需求复杂度的显示格式
                    modified_text = f"{modified_text}\n（实施复杂度：{complexity_val}）"
            rec[requirement_col] = modified_text

        for word_col, expr in compose_map.items():
            if word_col not in word_col_index:
                continue
            rec[word_col] = compose_row_value(
                expr,
                row,
                title_map=title_map,
                accept_word_col=accept_word_col,
                excel_columns=excel_columns,
            )

        new_rows.append(rec)

    if missing_targets:
        print("警告: Title_Map 中的下列 Word 列在模板表中不存在，已忽略这些映射：")
        for t in sorted(missing_targets):
            print(f" - {t}")

    if missing_compose_targets:
        print("警告: Compose_Map 中的下列 Word 列在模板表中不存在，已忽略这些规则：")
        for t in sorted(missing_compose_targets):
            print(f" - {t}")

    # 排序字段
    sort_col = accept_word_col or "需求审批\n时间"

    def sort_key(rec: dict[str, str]) -> tuple:
        d = parse_sort_date(rec.get(sort_col, ""))
        return (d or datetime.max, rec.get(sort_col, ""))

    new_rows.sort(key=sort_key)

    output_path = build_output_path(template_path)

    bookmark_id = str(abs(uuid.uuid4().int) % 900000000 + 100000000)
    first_new_row = True
    for rec in new_rows:
        row = append_cloned_row(table, template_tr)
        for ci, h in enumerate(header_cells):
            cell = row.cells[ci]
            text = rec.get(h, "")
            if first_new_row and ci == 0:
                set_cell_text_preserve_style_with_bookmark(
                    cell,
                    text,
                    bookmark_name=BOOKMARK_NEW_DATA,
                    bookmark_id=bookmark_id,
                )
                first_new_row = False
            else:
                set_cell_text_preserve_style(cell, text)

    added_count = len(new_rows)
    log_path = build_log_path(script_path)
    log_lines = [
        f"运行时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"配置文件: {config_path}",
        f"模板文件: {template_path}",
        f"使用表格索引: {table_index}",
        "输入文件:",
    ]
    log_lines.extend(f" - {p}" for p in input_paths)
    if missing_inputs:
        log_lines.append("缺失输入(已跳过):")
        log_lines.extend(f" - {p}" for p in missing_inputs)
    log_lines.extend(
        [
            f"输出文件: {output_path}",
            f"日志文件: {log_path}",
            f"模板原有表格数据行数: {original_data_rows}",
            f"本次新增记录数: {added_count}",
            f"输出表格总行数(含表头): {len(table.rows)}",
            "定位说明: Word 标准 docx 无法在打开时自动将光标跳到表格；"
            f"请使用「查找」->「转到」->「书签」-> 名称「{BOOKMARK_NEW_DATA}」跳转到首条新增行。",
        ]
    )
    if row_limit is not None:
        log_lines.append(f"INPUT_ROW_LIMIT 生效: 仅处理合并后的前 {row_limit} 行")

    doc.save(str(output_path))
    emit_log(log_lines, log_path)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print("程序执行失败：", file=sys.stderr)
        print(str(exc), file=sys.stderr)
        sys.exit(1)
