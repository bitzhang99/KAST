# -*- coding: utf-8 -*-
"""
双周报：从 Excel 台账合并需求行到 Word 模板表格（可配置 ini、多文件、日志、书签定位）。

程序修订版本号：V11
修订说明：新增行若未配置 `关联项目组` / `上线时间` 的映射关系，不再继承模板最后一条记录内容，而是强制写空。

需求修订编号（对应代码中 # REQ-xxx 注释）：
REQ-001 使用 openpyxl / python-docx 完成 Excel、Word 读写。
REQ-002 PyInstaller 打包时从 exe 所在目录读取 workreport.ini。
REQ-003 使用 utf-8-sig 读取 ini，兼容记事本 UTF-8 BOM。
REQ-004 新增行按“受理时间”升序；写入与排序均按 YYYY-MM-DD（可解析时）。
REQ-005 在首条新增行插入书签，便于打开文档后跳转到新数据。
REQ-006 通过 deepcopy 表格行 XML 追加行，保持与模板行格式一致。
REQ-007 Title_Map、Fixed_Columns 由配置文件驱动列映射与固定列。
REQ-008 支持 INPUT_PATHS 多个 Excel 合并后再排序写入。
REQ-009 运行前检查模板与各输入路径是否存在并给出可读错误说明。
REQ-010 日志写入“脚本名_日期时间.log”，并同步打印到标准输出。
REQ-011 输出 Word 文件名使用“模板主名_HHMMdd.docx”（时、分、日各两位，与 HHMMDD 书写一致；strftime 为 %H%M%d）。
"""
from __future__ import annotations

import configparser
import logging
import os
import re
import sys
from copy import deepcopy
from datetime import date, datetime
from typing import Any, Dict, List, Optional, Tuple

# -----------------------------------------------------------------------------
# REQ-001：依赖 openpyxl / python-docx 读取 Excel 与 Word
# -----------------------------------------------------------------------------
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from openpyxl import load_workbook

APP_VERSION = "V11"

# -----------------------------------------------------------------------------
# REQ-002：PyInstaller 打包后从可执行文件所在目录解析 workreport.ini
# -----------------------------------------------------------------------------
def _app_base_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))


# -----------------------------------------------------------------------------
# REQ-003：兼容 Windows 记事本“UTF-8 带 BOM”保存方式，避免 ConfigParser 读取出错
# -----------------------------------------------------------------------------
def _load_ini(path: str) -> configparser.ConfigParser:
    cfg = configparser.ConfigParser(
        interpolation=None,
        inline_comment_prefixes=("#", ";"),
    )
    cfg.optionxform = str  # 保留键名大小写（Title_Map 中文列名）
    with open(path, "r", encoding="utf-8-sig", newline=None) as f:
        cfg.read_file(f)
    return cfg


def _split_input_paths(raw: str) -> List[str]:
    if not raw or not str(raw).strip():
        return []
    parts = re.split(r"[,，\n\r]+", raw)
    return [p.strip().strip('"').strip("'") for p in parts if p.strip()]


def _normalize_header(s: Any) -> str:
    if s is None:
        return ""
    t = str(s).replace("\u3000", " ").strip()
    t = re.sub(r"\s+", "", t)
    return t


def _parse_sort_date(raw: Any) -> date:
    """REQ-004：按受理时间 YYYY-MM-DD 排序；无法解析的排到最后。"""
    if raw is None or (isinstance(raw, str) and not raw.strip()):
        return date(9999, 12, 31)
    if isinstance(raw, datetime):
        return raw.date()
    if isinstance(raw, date):
        return raw
    s = str(raw).strip()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    m = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", s)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return date(y, mo, d)
        except ValueError:
            pass
    return date(9999, 12, 31)


def _format_accept_date(raw: Any) -> str:
    """REQ-004：写入 Word 时统一为 YYYY-MM-DD（能解析则格式化）。"""
    d = _parse_sort_date(raw)
    if d == date(9999, 12, 31):
        return "" if raw is None else str(raw).strip()
    return d.strftime("%Y-%m-%d")


def _next_bookmark_id(document: DocumentObject) -> str:
    """REQ-005：插入书签前分配未占用的 w:id。"""
    max_id = 0
    root = document.element.body
    for el in root.iter(qn("w:bookmarkStart")):
        bid = el.get(qn("w:id"))
        if bid and bid.isdigit():
            max_id = max(max_id, int(bid))
    return str(max_id + 1)


def _insert_bookmark_on_paragraph(paragraph, name: str, bookmark_id: str) -> None:
    """REQ-005：打开 Word 后可通过“插入-链接-书签”或 Ctrl+G 转到定位。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    p = paragraph._p
    p.insert(0, start)
    p.append(end)


def _set_cell_text_keep_style(cell, text: str) -> None:
    """REQ-006：克隆行后仅改文字，尽量保留首个 run 的字体/段落样式。"""
    text = "" if text is None else str(text)
    tc = cell._tc
    for el in tc.findall(qn("w:p"))[1:]:
        tc.remove(el)
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p0 = cell.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(text)


def _find_target_table(document: DocumentObject, required_word_headers: List[str]) -> Tuple[Table, Dict[str, int], int]:
    """
    REQ-007：定位含所需表头的表格，并返回列索引；header_row_index 为表头所在行号。
    """
    req_norm = {_normalize_header(h): h for h in required_word_headers}
    best: Optional[Tuple[Table, Dict[str, int], int]] = None
    best_score = -1

    for tbl in document.tables:
        for ri, row in enumerate(tbl.rows):
            texts = [_normalize_header(c.text) for c in row.cells]
            col_map: Dict[str, int] = {}
            for ci, key in enumerate(texts):
                if key in req_norm:
                    word_h = req_norm[key]
                    col_map[word_h] = ci
            score = len(col_map)
            if score > best_score:
                best_score = score
                best = (tbl, col_map, ri)

    if best is None or best_score < len(required_word_headers):
        missing = set(required_word_headers)
        if best:
            missing -= set(best[1].keys())
        raise ValueError(
            "未在 Word 模板中找到包含全部目标列的表格。"
            f" 需要列: {required_word_headers}；缺失: {sorted(missing)}"
        )
    return best  # type: ignore[return-value]


def _get_header_col_indices(table: Table, header_row_index: int, target_headers: List[str]) -> Dict[str, int]:
    """
    在已定位的表格中查找“表头行”（可能不止一行），找到目标表头列对应的列索引。
    用于可选列处理：当配置未给出映射关系时，也可以强制覆盖这两列写空。
    """
    norm_targets = {_normalize_header(h): h for h in target_headers}
    best_row_index = header_row_index
    best_score = 0

    # 遍历表格行，找“匹配目标表头最多”的那行，避免模板表头分成多行导致找错行。
    for ri, row in enumerate(table.rows):
        score = 0
        for cell in row.cells:
            if _normalize_header(cell.text) in norm_targets:
                score += 1
        if score > best_score:
            best_score = score
            best_row_index = ri
        if best_score >= len(target_headers):
            break

    if best_score <= 0:
        return {}

    out: Dict[str, int] = {}
    header_row = table.rows[best_row_index]
    for ci, cell in enumerate(header_row.cells):
        c_norm = _normalize_header(cell.text)
        if c_norm in norm_targets:
            out[norm_targets[c_norm]] = ci
    return out


def _read_excel_rows(
    path: str,
    title_map: Dict[str, str],
    fixed_cols: Dict[str, str],
    max_rows: int,
) -> Tuple[List[Dict[str, str]], Dict[str, int]]:
    """REQ-008：从单个 Excel 读取数据行；返回行字典（键为 Word 列名）与列索引映射。"""
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        header_row = next(rows_iter, None)
        if not header_row:
            return [], {}
        excel_to_idx: Dict[str, int] = {}
        for i, h in enumerate(header_row):
            nh = _normalize_header(h)
            if nh:
                excel_to_idx[nh] = i

        out: List[Dict[str, str]] = []
        count = 0
        for row in rows_iter:
            if row is None or all(v is None or str(v).strip() == "" for v in row):
                continue
            rec: Dict[str, str] = {}
            for excel_name, word_name in title_map.items():
                key = _normalize_header(excel_name)
                if key not in excel_to_idx:
                    continue
                idx = excel_to_idx[key]
                val = row[idx] if idx < len(row) else None
                if word_name == "受理时间":
                    rec[word_name] = _format_accept_date(val)
                else:
                    rec[word_name] = "" if val is None else str(val).strip()
            for wcol, fval in fixed_cols.items():
                rec[wcol] = fval
            out.append(rec)
            count += 1
            if max_rows > 0 and count >= max_rows:
                break
        return out, excel_to_idx
    finally:
        wb.close()


def _append_cloned_rows(table: Table, n: int, style_row_index: int) -> List[int]:
    """REQ-006：追加 n 行，克隆模板行 XML 以保留表格样式。"""
    tbl = table._tbl
    tr = table.rows[style_row_index]._tr
    new_indices: List[int] = []
    for _ in range(n):
        new_tr = deepcopy(tr)
        tbl.append(new_tr)
        new_indices.append(len(table.rows) - 1)
    return new_indices


def _validate_paths(template_path: str, input_paths: List[str]) -> None:
    """REQ-009：路径存在性预检查与友好报错。"""
    errs: List[str] = []
    if not template_path:
        errs.append("TEMPLATE_PATH 为空，请在 workreport.ini 的 [Paths] 中配置 Word 模板路径。")
    elif not os.path.isfile(template_path):
        errs.append(
            f"模板文件不存在或不是文件: {template_path}\n"
            f"  建议：检查盘符、目录名是否一致；是否在资源管理器中确认该文件存在。"
        )
    if not input_paths:
        errs.append(
            "INPUT_PATHS 未配置或为空。\n"
            "  建议：在 workreport.ini 的 [Paths] 中写入一个或多个 Excel 完整路径，英文逗号分隔。"
        )
    for p in input_paths:
        if not p:
            continue
        if not os.path.isfile(p):
            errs.append(
                f"Excel 输入文件不存在或不是文件: {p}\n"
                f"  建议：核对 INPUT_PATHS 是否含多余逗号、是否使用了全角逗号（可混用但路径须正确）。"
            )
    if errs:
        raise FileNotFoundError("\n".join(errs))


def _setup_logging(script_stem: str, base_dir: str) -> Tuple[logging.Logger, str]:
    """REQ-010：日志文件名 = 脚本名 + 运行日期时间 + .log，并同时输出到控制台。"""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = os.path.join(base_dir, f"{script_stem}_{ts}.log")
    logger = logging.getLogger("biweekly_word")
    logger.setLevel(logging.INFO)
    logger.handlers.clear()
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler(sys.stdout)
    sh.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(sh)
    return logger, log_path


def main() -> int:
    base_dir = _app_base_dir()
    ini_path = os.path.join(base_dir, "workreport.ini")

    script_stem = os.path.splitext(os.path.basename(sys.argv[0] if sys.argv else __file__))[0]
    logger, log_path = _setup_logging(script_stem, base_dir)

    try:
        if not os.path.isfile(ini_path):
            raise FileNotFoundError(
                f"未找到配置文件: {ini_path}\n"
                f"  建议：将 workreport.ini 放在与本程序同一目录（打包后放在 exe 同目录）。"
            )

        cfg = _load_ini(ini_path)  # REQ-003

        template_path = cfg.get("Paths", "TEMPLATE_PATH", fallback="").strip()
        input_paths = _split_input_paths(cfg.get("Paths", "INPUT_PATHS", fallback=""))
        max_rows = int(cfg.get("Settings", "MAX_IMPORT_ROWS", fallback="0") or "0")
        bookmark_name = cfg.get("Settings", "NEW_DATA_BOOKMARK", fallback="NewImportedRows").strip() or "NewImportedRows"

        title_map: Dict[str, str] = {}
        if cfg.has_section("Title_Map"):
            for k, v in cfg.items("Title_Map"):
                if k.startswith("#") or k.startswith(";"):
                    continue
                kk, vv = k.strip(), v.split("#")[0].split(";")[0].strip()
                if kk and vv:
                    title_map[kk] = vv

        fixed_cols: Dict[str, str] = {}
        if cfg.has_section("Fixed_Columns"):
            for k, v in cfg.items("Fixed_Columns"):
                kk, vv = k.strip(), v.split("#")[0].split(";")[0].strip()
                if kk:
                    fixed_cols[kk] = vv

        _validate_paths(template_path, input_paths)  # REQ-009

        word_headers = list(dict.fromkeys(list(title_map.values()) + list(fixed_cols.keys())))

        all_records: List[Tuple[date, Dict[str, str]]] = []
        for xlsx in input_paths:
            rows, _hdr_idx = _read_excel_rows(xlsx, title_map, fixed_cols, max_rows)  # REQ-008
            for r in rows:
                sort_key = _parse_sort_date(r.get("受理时间", ""))  # REQ-004
                all_records.append((sort_key, r))

        all_records.sort(key=lambda x: (x[0], str(x[1])))

        doc = Document(template_path)
        table, col_map, header_row_index = _find_target_table(doc, word_headers)  # REQ-007

        # REQ-012：当 `关联项目组` / `上线时间` 没有映射配置时，新行会继承“克隆行”的原内容；
        # 因而这里对这两列做强制覆盖：映射不存在则写空。
        optional_headers = ["关联项目组", "上线时间"]
        optional_col_map = _get_header_col_indices(table, header_row_index, optional_headers)

        # REQ-006：与「前面已有表格行」格式一致——优先克隆末行数据行，而非表头后第一行（首行常为占位/样式不同）
        last_idx = len(table.rows) - 1
        style_row_index = last_idx if last_idx > header_row_index else header_row_index
        if style_row_index == header_row_index:
            logger.warning(
                "模板表格可能仅有表头行，新行格式将克隆表头；建议在模板中保留至少一行正文数据行以匹配既有表格风格。"
            )

        n_new = len(all_records)
        if n_new > 0:
            new_row_indices = _append_cloned_rows(table, n_new, style_row_index)  # REQ-006
            for idx, (_, rec) in enumerate(all_records):
                row = table.rows[new_row_indices[idx]]

                # 归一化 key：避免配置/模板存在全角空格等细微差异导致 rec.get() 命中失败
                rec_norm = {_normalize_header(k): v for k, v in rec.items()}

                for wcol, ci in col_map.items():
                    if ci < len(row.cells):
                        _set_cell_text_keep_style(row.cells[ci], rec.get(wcol, ""))

                for opt_wcol, ci in optional_col_map.items():
                    if ci < len(row.cells):
                        _set_cell_text_keep_style(
                            row.cells[ci],
                            rec_norm.get(_normalize_header(opt_wcol), ""),
                        )

            first_new = table.rows[new_row_indices[0]]
            bid = _next_bookmark_id(doc)  # REQ-005
            if first_new.cells and first_new.cells[0].paragraphs:
                _insert_bookmark_on_paragraph(first_new.cells[0].paragraphs[0], bookmark_name, bid)

        template_dir = os.path.dirname(os.path.abspath(template_path))
        template_stem = os.path.splitext(os.path.basename(template_path))[0]
        out_ts = datetime.now().strftime("%Y%m%d%H%M%S")
        out_name = f"{template_stem}_{out_ts}.docx"  # REQ-011：HHMMdd（时、分、日）
        out_path = os.path.join(template_dir, out_name)

        doc.save(out_path)

        logger.info("配置文件: %s", ini_path)
        logger.info("模板: %s", template_path)
        logger.info("程序版本: %s", APP_VERSION)
        logger.info("输入 Excel 数量: %d", len(input_paths))
        logger.info("本次新增表格行数: %d", n_new)
        logger.info("输出 Word: %s", out_path)
        logger.info("日志文件: %s", log_path)
        if n_new > 0:
            logger.info('书签名称: %s（在 Word 中按 Ctrl+G / “转到” 选择书签）', bookmark_name)

        print(f"[完成] {APP_VERSION} 已写入: {out_path}", flush=True)
        return 0

    except Exception as e:
        logger.exception("运行失败: %s", e)
        print(f"[错误] {e}", file=sys.stderr, flush=True)
        return 1


if __name__ == "__main__":
    sys.exit(main())
